import os

from PyQt5.QtWidgets import QWidget, QTreeWidgetItem
from docx import Document

from widgets.dictSelect.widget_manual import widget_manual


class Manual(QWidget, widget_manual):
    def __init__(self, parent=None):
        super(Manual, self).__init__(parent)
        self.setupUi(self)
        self.initWord()

    def initWord(self):
        self.headingTxt = {}
        self.allTxt = {}
        self.tabNum = []
        path = os.path.abspath('装备性能手册.docx')
        self.wordFile = Document(path)
        for i, p in enumerate(self.wordFile.paragraphs):
            style_name = p.style.name
            self.allTxt[i] = p.text
            # print(p.row,p.text,sep=':')
            if style_name.startswith('Heading'):
                self.headingTxt[i] = p.text
                self.tabNum.append(int(style_name.lstrip('Heading ')))
                # print(style_name, p.text, sep=':')
        # print("self.allTxt",self.allTxt,"\nself.headingTxt = ",self.headingTxt)
        self.initHeading()
        self.initArticle()

        # for i in self.headingTxt:
        #     if i
        # stack = self.headingTxt
        # root = []
        # for i in self.headingTxt.values():
        #     stack.append(i)
        # root.append(self.tw_catalog)
        # self.initHeading(stack,root)

    def initHeading(self):
        minTab = min(self.tabNum)
        for i, xx in enumerate(self.headingTxt.values()):
            count = self.tabNum[i] - minTab
            txt = count * '   ' + xx
            item = QTreeWidgetItem()
            item.setText(0, txt)
            self.tw_catalog.addTopLevelItem(item)

    # def initHeading(self,stack,root):
    #     while stack:
    #         EquipInfo = stack.pop(0)
    #         item = QTreeWidgetItem(root.pop(0))
    #         item.setText(0, EquipInfo)
    #         # self.second_treeWidget_dict[EquipInfo[0]] = item
    #         root.append(item)
    #         # result = selectEquipInfoByEquipUper(EquipInfo[0])
    #         # for resultInfo in result:
    #         #     stack.append(resultInfo)
    #         #     root.append(item)

    def initArticle(self):
        # self.tb_article.setProperty("Color",QColor(255,250,250))
        for p in self.wordFile.paragraphs:
            style_name = p.style.name
            self.tb_article.setStyleSheet('background-color:white;'
                                          'font-size:11px;'
                                          'color:black')
            if style_name.startswith('Title'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:24pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Heading 1'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:22pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Heading 2'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:20pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Heading 3'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:18pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Heading 4'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:16pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Heading 5'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:14pt; font-weight:600;">%s</span></p>' % p.text)
            elif style_name.startswith('Normal'):
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:11pt;">%s</span></p>' % p.text)
            else:
                self.tb_article.append('<p align="left" style=" margin-top:0px; margin-bottom:0px; margin-left:0px; '
                                       'margin-right:0px; -qt-block-indent:0; text-indent:0px;"><span style=" '
                                       'font-size:12pt;font-weight:600;">%s</span></p>' % p.text)

    # def initArticle(self):
    #     self.axWidget.clear()
    #     # self.closeWord()
    #     # self.axWidget = QAxWidget("Word.Application",self.widget)
    #     if not self.axWidget.setControl('Word.Application'):
    #         return QMessageBox.critical(self, '错误', '%s有误' %'Word.Application')
    #     self.horizontalLayout.addWidget(self.axWidget)
    #     self.axWidget.dynamicCall('SetVisible (bool Visible)', 'false')
    #     self.axWidget.setProperty("DisplayAlerts", False)
    #     rect = self.widget.geometry()
    #     self.axWidget.setGeometry(rect)
    #     path1 = os.path.abspath('装备性能手册.docx')
    #     self.axWidget.setControl(path1)
    #     self.axWidget.show()
    #

    # def initArticle(self):
    #     self.axWidget.clear()
    #     path = os.path.abspath('装备性能手册.docx')
    #     if not self.axWidget.setControl('Kwps.Application'):
    #         return QMessageBox.critical(self, '错误', '没有安装  %s' % 'Kwps.Application')
    #     self.axWidget.dynamicCall(
    #         'SetVisible (bool Visible)', 'false')  # 不显示窗体
    #     self.axWidget.setProperty('DisplayAlerts', False)
    #     self.axWidget.setControl(path)
    #     self.axWidget.show()

    # def closeEvent(self, event):
    #     self.axWidget.close()
    #     self.axWidget.clear()
    #     self.layout().removeWidget(self.axWidget)
    #     del self.axWidget
    #     super(Manual, self).closeEvent(event)
